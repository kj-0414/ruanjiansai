import os
import json
import shutil
import re
from datetime import datetime
from modules.job.job_repository import JobRepository
from modules.job.schemas import JobCreate, JobUpdate, JobAbilityCreate
from modules.job.exceptions import JobNotFoundException, AccessDeniedException
from modules.agents.job_agent import JobParseAgent
from modules.agents.graph_agent import GraphAgent
from sqlalchemy.orm import Session
from fastapi import UploadFile


# 全局Agent实例
_job_agent = JobParseAgent()
_graph_agent = GraphAgent()


def get_knowledge_base_keywords():
    """从知识库获取技能关键词（增强版）"""
    try:
        from modules.knowledge_base import KnowledgeBaseService
        kb = KnowledgeBaseService()
        skills = kb.db.get_all_skills()

        keywords = {}
        for skill in skills:
            category = skill.get('category', '其他')
            if category not in keywords:
                keywords[category] = []

            skill_name = skill.get('skill_name', '')
            keywords[category].append(skill_name.lower())

            aliases_text = skill.get('skill_aliases', '[]')
            try:
                aliases = json.loads(aliases_text)
                for alias in aliases:
                    if alias.lower() not in keywords[category]:
                        keywords[category].append(alias.lower())
            except:
                pass

        kb.close()
        return keywords
    except Exception:
        return None


def clean_special_chars(text):
    if not text:
        return text
    try:
        text.encode('gbk')
        return text
    except UnicodeEncodeError:
        cleaned_text = ''
        for char in text:
            try:
                char.encode('gbk')
                cleaned_text += char
            except UnicodeEncodeError:
                cleaned_text += ' '
        return cleaned_text


def extract_job_info_from_text(text: str, filename: str) -> dict:
    """使用Agent解析岗位信息（fallback到正则解析）"""
    try:
        # 使用Agent解析
        result = _job_agent.parse_job(text)
        
        return {
            "jobName": result.get("job_title", re.sub(r'\.[^/.]+$', "", filename) if isinstance(filename, str) else "未知岗位"),
            "salary": result.get("salary_range", ""),
            "location": result.get("location", ""),
            "education": result.get("required_education", "不限"),
            "experience": result.get("required_experience", ""),
            "recruitmentCount": 1,
            "description": "\n".join(result.get("job_responsibilities", [])),
            "responsibilities": "\n".join(result.get("job_responsibilities", [])),
            "requirements": "\n".join(result.get("job_requirements", [])),
            "skills": result.get("required_skills", []),
            "_agent_result": result  # 保留原始Agent解析结果
        }
    except Exception as e:
        print(f"[JobService] Agent解析失败: {e}")
        # Fallback到正则解析
        text_lower = text.lower()
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        result = {
            "jobName": re.sub(r'\.[^/.]+$', "", filename) if isinstance(filename, str) else "未知岗位",
            "salary": "",
            "location": "",
            "education": "不限",
            "experience": "",
            "recruitmentCount": 1,
            "description": "",
            "responsibilities": "",
            "requirements": "",
            "skills": []
        }

        job_name_patterns = [
            r'(岗位名称|职位名称|job title|position)\s*[：:]\s*([^\n]+)',
            r'(招聘|诚聘)\s*([^\n]+)',
            r'【([^】]+)】',
            r'《([^》]+)》'
        ]
        for pattern in job_name_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                candidate = match.group(2).strip() if len(match.groups()) > 1 else match.group(1).strip()
                if candidate and len(candidate) > 2 and len(candidate) < 50:
                    result["jobName"] = candidate
                    break

        salary_patterns = [
            r'薪资\s*[：:]\s*([^\n]+)',
            r'salary\s*[：:]\s*([^\n]+)',
            r'(\d+)[-~至](\d+)\s*万',
            r'(\d+)K[-~至](\d+)K',
            r'(\d+)K以上',
            r'(\d+)K\+'
        ]
        for pattern in salary_patterns:
            match = re.search(pattern, text)
            if match:
                if len(match.groups()) == 2:
                    result["salary"] = f"{match.group(1)}K-{match.group(2)}K"
                else:
                    salary_value = match.group(1).strip()
                    if '福利' not in salary_value and '五险' not in salary_value and '一金' not in salary_value:
                        result["salary"] = f"{salary_value}K以上"
                break

        location_patterns = [
            r'工作地点\s*[：:]\s*([^\n]+)',
            r'地点\s*[：:]\s*([^\n]+)',
            r'location\s*[：:]\s*([^\n]+)',
            r'(省|市|区)\s*([^\n]+)'
        ]
        for pattern in location_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                loc = match.group(1).strip()
                if len(loc) > 2:
                    result["location"] = loc
                break

        education_keywords = {
            '博士': ['博士', 'phd', '博士后'],
            '硕士': ['硕士', '研究生', '硕士研究生', 'master'],
            '本科': ['本科', '学士', 'bachelor'],
            '大专': ['大专', '专科']
        }
        for level, keywords in education_keywords.items():
            for keyword in keywords:
                if keyword.lower() in text_lower:
                    context_start = max(0, text_lower.find(keyword) - 20)
                    context_end = min(len(text_lower), text_lower.find(keyword) + 20)
                    context = text_lower[context_start:context_end]
                    if '学历' in context or '学位' in context or '要求' in context or '本科及以上' in text_lower:
                        result["education"] = level
                        break
            if result["education"] != "不限":
                break

        experience_patterns = [
            r'经验\s*[：:]\s*([^\n]+)',
            r'工作经验\s*[：:]\s*([^\n]+)',
            r'experience\s*[：:]\s*([^\n]+)',
            r'(\d+)[-~至](\d+)\s*年',
            r'(\d+)\s*年以上',
            r'(\d+)\s*年.*经验'
        ]
        for pattern in experience_patterns:
            match = re.search(pattern, text)
            if match:
                if len(match.groups()) == 2:
                    result["experience"] = f"{match.group(1)}-{match.group(2)}年"
                else:
                    exp_value = match.group(1).strip()
                    if '年' not in exp_value:
                        exp_value += '年'
                    result["experience"] = exp_value
                break

        count_patterns = [
            r'招聘人数\s*[：:]\s*(\d+)',
            r'人数\s*[：:]\s*(\d+)',
            r'招\s*(\d+)人'
        ]
        for pattern in count_patterns:
            match = re.search(pattern, text)
            if match:
                result["recruitmentCount"] = int(match.group(1))
                break

        responsibilities_keywords = ['职责', '岗位职责', '工作职责', '主要职责', 'responsibilities', 'duties']
        responsibilities_lines = []
        in_responsibilities = False
        for i, line in enumerate(lines):
            line_lower = line.lower()
            if any(kw in line_lower for kw in responsibilities_keywords):
                in_responsibilities = True
                continue
            if in_responsibilities:
                exit_keywords = ['要求', '任职要求', '资格', '福利', '待遇', '薪资', '描述']
                if any(kw in line_lower for kw in exit_keywords) and len(line) < 30:
                    break
                if line and not line.startswith('●') and not line.startswith('◆') and not line.startswith('■'):
                    responsibilities_lines.append(line)
                    if len(responsibilities_lines) >= 8:
                        break
        if responsibilities_lines:
            result["responsibilities"] = '\n'.join(responsibilities_lines)

        requirements_keywords = ['要求', '任职要求', '资格要求', '岗位要求', 'qualifications', 'requirements']
        requirements_lines = []
        in_requirements = False
        found_break_keyword = False
        for i, line in enumerate(lines):
            line_lower = line.lower()
            if found_break_keyword:
                continue
            if any(kw in line_lower for kw in requirements_keywords):
                in_requirements = True
                continue
            if in_requirements and ('福利' in line_lower or '其他信息' in line_lower):
                found_break_keyword = True
                continue
            if in_requirements:
                if len(line) < 10 and any(kw in line_lower for kw in ['职责', '描述', '说明']):
                    continue
                if line:
                    requirements_lines.append(line)
                    if len(requirements_lines) >= 15:
                        break
        if requirements_lines:
            result["requirements"] = '\n'.join(requirements_lines)

        desc_keywords = ['描述', '岗位描述', '职位描述', 'job description', '简介']
        desc_lines = []
        in_desc = False
        for i, line in enumerate(lines):
            line_lower = line.lower()
            if any(kw in line_lower for kw in desc_keywords):
                in_desc = True
                continue
            if in_desc:
                exit_keywords = ['职责', '要求', '任职']
                if any(kw in line_lower for kw in exit_keywords) and len(line) < 30:
                    break
                if line and len(line) > 10:
                    desc_lines.append(line)
                    if len(desc_lines) >= 5:
                        break
        if desc_lines:
            result["description"] = '\n'.join(desc_lines)
        elif not result["description"]:
            result["description"] = text[:500]

        kb_keywords = get_knowledge_base_keywords()
        if kb_keywords:
            skill_keywords = kb_keywords
        else:
            skill_keywords = {
                '编程语言': ['python', 'java', 'javascript', 'typescript', 'go', 'rust', 'c++', 'c#', 'c语言', 'php', 'ruby', 'scala', 'kotlin', 'swift', 'objective-c', 'dart', 'lua'],
                '前端技术': ['vue', 'react', 'angular', 'svelte', 'next.js', 'nuxt.js', 'html', 'css', 'sass', 'less', 'webpack', 'vite', 'babel', 'jquery', 'bootstrap', 'tailwind'],
                '后端框架': ['spring', 'django', 'flask', 'fastapi', 'express', 'nestjs', 'egg.js', 'koa', 'laravel', 'symfony', 'rails', 'gin', 'echo'],
                '数据库': ['sql', 'mysql', 'postgresql', 'oracle', 'sqlserver', 'mongodb', 'redis', 'elasticsearch', 'cassandra', 'sqlite', 'hbase', '达梦', 'dm'],
                'DevOps': ['docker', 'kubernetes', 'k8s', 'jenkins', 'gitlab-ci', 'github-actions', 'ansible', 'terraform', 'prometheus', 'grafana', 'elk', 'nginx', 'apache'],
                '云计算': ['aws', 'azure', 'gcp', '阿里云', '腾讯云', '华为云', 'ecs', 's3', 'lambda', 'serverless'],
                'AI/机器学习': ['tensorflow', 'pytorch', 'keras', 'scikit-learn', 'nlp', '计算机视觉', 'cv', '深度学习', '机器学习', '人工智能', 'ai', '大语言模型', 'llm'],
                '移动开发': ['android', 'ios', 'react-native', 'flutter', 'uniapp'],
                '其他': ['git', 'linux', 'unix', 'shell', 'bash', '算法', '数据结构', '微服务', 'grpc', 'restful', 'graphql', '消息队列', 'rabbitmq', 'kafka', 'rocketmq']
            }

        found_skills = []
        for category, keywords in skill_keywords.items():
            for keyword in keywords:
                if keyword.lower() in text_lower:
                    found_skills.append(keyword)

        if not found_skills:
            found_skills = ['Python', 'Java', 'SQL', 'Git']

        result["skills"] = list(set(found_skills))[:12]

        return result

    result = {
        "jobName": re.sub(r'\.[^/.]+$', "", filename) if isinstance(filename, str) else "未知岗位",
        "salary": "",
        "location": "",
        "education": "不限",
        "experience": "",
        "recruitmentCount": 1,
        "description": "",
        "responsibilities": "",
        "requirements": "",
        "skills": []
    }

    job_name_patterns = [
        r'(岗位名称|职位名称|job title|position)\s*[：:]\s*([^\n]+)',
        r'(招聘|诚聘)\s*([^\n]+)',
        r'【([^】]+)】',
        r'《([^》]+)》'
    ]
    for pattern in job_name_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            candidate = match.group(2).strip() if len(match.groups()) > 1 else match.group(1).strip()
            if candidate and len(candidate) > 2 and len(candidate) < 50:
                result["jobName"] = candidate
                break

    salary_patterns = [
        r'薪资\s*[：:]\s*([^\n]+)',
        r'salary\s*[：:]\s*([^\n]+)',
        r'(\d+)[-~至](\d+)\s*万',
        r'(\d+)K[-~至](\d+)K',
        r'(\d+)K以上',
        r'(\d+)K\+'
    ]
    for pattern in salary_patterns:
        match = re.search(pattern, text)
        if match:
            if len(match.groups()) == 2:
                result["salary"] = f"{match.group(1)}K-{match.group(2)}K"
            else:
                salary_value = match.group(1).strip()
                if '福利' not in salary_value and '五险' not in salary_value and '一金' not in salary_value:
                    result["salary"] = f"{salary_value}K以上"
            break

    location_patterns = [
        r'工作地点\s*[：:]\s*([^\n]+)',
        r'地点\s*[：:]\s*([^\n]+)',
        r'location\s*[：:]\s*([^\n]+)',
        r'(省|市|区)\s*([^\n]+)'
    ]
    for pattern in location_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            loc = match.group(1).strip()
            if len(loc) > 2:
                result["location"] = loc
            break

    education_keywords = {
        '博士': ['博士', 'phd', '博士后'],
        '硕士': ['硕士', '研究生', '硕士研究生', 'master'],
        '本科': ['本科', '学士', 'bachelor'],
        '大专': ['大专', '专科']
    }
    for level, keywords in education_keywords.items():
        for keyword in keywords:
            if keyword.lower() in text_lower:
                context_start = max(0, text_lower.find(keyword) - 20)
                context_end = min(len(text_lower), text_lower.find(keyword) + 20)
                context = text_lower[context_start:context_end]
                if '学历' in context or '学位' in context or '要求' in context or '本科及以上' in text_lower:
                    result["education"] = level
                    break
        if result["education"] != "不限":
            break

    experience_patterns = [
        r'经验\s*[：:]\s*([^\n]+)',
        r'工作经验\s*[：:]\s*([^\n]+)',
        r'experience\s*[：:]\s*([^\n]+)',
        r'(\d+)[-~至](\d+)\s*年',
        r'(\d+)\s*年以上',
        r'(\d+)\s*年.*经验'
    ]
    for pattern in experience_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            if len(match.groups()) == 2:
                result["experience"] = f"{match.group(1)}-{match.group(2)}年"
            else:
                exp_value = match.group(1).strip()
                if '年' not in exp_value:
                    exp_value += '年'
                result["experience"] = exp_value
            break

    count_patterns = [
        r'招聘人数\s*[：:]\s*(\d+)',
        r'人数\s*[：:]\s*(\d+)',
        r'招\s*(\d+)人'
    ]
    for pattern in count_patterns:
        match = re.search(pattern, text)
        if match:
            result["recruitmentCount"] = int(match.group(1))
            break

    responsibilities_keywords = ['职责', '岗位职责', '工作职责', '主要职责', 'responsibilities', 'duties']
    responsibilities_lines = []
    in_responsibilities = False
    for i, line in enumerate(lines):
        line_lower = line.lower()
        if any(kw in line_lower for kw in responsibilities_keywords):
            in_responsibilities = True
            continue
        if in_responsibilities:
            exit_keywords = ['要求', '任职要求', '资格', '福利', '待遇', '薪资', '描述']
            if any(kw in line_lower for kw in exit_keywords) and len(line) < 30:
                break
            if line and not line.startswith('●') and not line.startswith('◆') and not line.startswith('■'):
                responsibilities_lines.append(line)
                if len(responsibilities_lines) >= 8:
                    break
    if responsibilities_lines:
        result["responsibilities"] = '\n'.join(responsibilities_lines)

    requirements_keywords = ['要求', '任职要求', '资格要求', '岗位要求', 'qualifications', 'requirements']
    requirements_lines = []
    in_requirements = False
    found_break_keyword = False
    for i, line in enumerate(lines):
        line_lower = line.lower()
        if found_break_keyword:
            continue
        if any(kw in line_lower for kw in requirements_keywords):
            in_requirements = True
            continue
        if in_requirements and ('福利' in line_lower or '其他信息' in line_lower):
            found_break_keyword = True
            continue
        if in_requirements:
            if len(line) < 10 and any(kw in line_lower for kw in ['职责', '描述', '说明']):
                continue
            if line:
                requirements_lines.append(line)
                if len(requirements_lines) >= 15:
                    break
    if requirements_lines:
        result["requirements"] = '\n'.join(requirements_lines)

    desc_keywords = ['描述', '岗位描述', '职位描述', 'job description', '简介']
    desc_lines = []
    in_desc = False
    for i, line in enumerate(lines):
        line_lower = line.lower()
        if any(kw in line_lower for kw in desc_keywords):
            in_desc = True
            continue
        if in_desc:
            exit_keywords = ['职责', '要求', '任职']
            if any(kw in line_lower for kw in exit_keywords) and len(line) < 30:
                break
            if line and len(line) > 10:
                desc_lines.append(line)
                if len(desc_lines) >= 5:
                    break
    if desc_lines:
        result["description"] = '\n'.join(desc_lines)
    elif not result["description"]:
        result["description"] = text[:500]

    skill_keywords = {
        '编程语言': ['python', 'java', 'javascript', 'typescript', 'go', 'rust', 'c++', 'c#', 'c语言', 'php', 'ruby', 'scala', 'kotlin', 'swift', 'objective-c', 'dart', 'lua'],
        '前端技术': ['vue', 'react', 'angular', 'svelte', 'next.js', 'nuxt.js', 'html', 'css', 'sass', 'less', 'webpack', 'vite', 'babel', 'jquery', 'bootstrap', 'tailwind'],
        '后端框架': ['spring', 'django', 'flask', 'fastapi', 'express', 'nestjs', 'egg.js', 'koa', 'laravel', 'symfony', 'rails', 'gin', 'echo'],
        '数据库': ['sql', 'mysql', 'postgresql', 'oracle', 'sqlserver', 'mongodb', 'redis', 'elasticsearch', 'cassandra', 'sqlite', 'hbase', '达梦', 'dm'],
        'DevOps': ['docker', 'kubernetes', 'k8s', 'jenkins', 'gitlab-ci', 'github-actions', 'ansible', 'terraform', 'prometheus', 'grafana', 'elk', 'nginx', 'apache'],
        '云计算': ['aws', 'azure', 'gcp', '阿里云', '腾讯云', '华为云', 'ecs', 's3', 'lambda', 'serverless'],
        'AI/机器学习': ['tensorflow', 'pytorch', 'keras', 'scikit-learn', 'nlp', '计算机视觉', 'cv', '深度学习', '机器学习', '人工智能', 'ai', '大语言模型', 'llm'],
        '移动开发': ['android', 'ios', 'react-native', 'flutter', 'uniapp'],
        '其他': ['git', 'linux', 'unix', 'shell', 'bash', '算法', '数据结构', '微服务', 'grpc', 'restful', 'graphql', '消息队列', 'rabbitmq', 'kafka', 'rocketmq']
    }

    found_skills = []
    for category, keywords in skill_keywords.items():
        for keyword in keywords:
            if keyword.lower() in text_lower:
                found_skills.append(keyword)

    if not found_skills:
        found_skills = ['Python', 'Java', 'SQL', 'Git']

    result["skills"] = list(set(found_skills))[:12]

    return result


class JobService:
    def __init__(self):
        self.repository = None

    def _init_repository(self, db: Session):
        if not self.repository:
            self.repository = JobRepository(db)

    def _build_job_response(self, job) -> dict:
        result = {
            "id": job.id,
            "job_name": job.job_name,
            "job_desc": job.job_desc,
            "salary": job.salary,
            "location": job.location,
            "work_hours": job.work_hours,
            "education_requirement": job.education_requirement,
            "experience_requirement": job.experience_requirement,
            "recruitment_count": job.recruitment_count,
            "department": job.department,
            "job_type": job.job_type,
            "benefits": job.benefits,
            "user_id": job.user_id,
            "create_time": job.created_at.strftime("%Y-%m-%d %H:%M:%S")
        }

        if hasattr(job, 'responsibilities'):
            result["responsibilities"] = job.responsibilities
        if hasattr(job, 'requirements'):
            result["requirements"] = job.requirements
        if hasattr(job, 'skills'):
            result["skills"] = job.skills
        if hasattr(job, 'skills_requirement'):
            result["skills_requirement"] = job.skills_requirement
        if hasattr(job, 'certificate_requirements'):
            result["certificate_requirements"] = job.certificate_requirements
        if hasattr(job, 'project_requirements'):
            result["project_requirements"] = job.project_requirements
        if hasattr(job, 'education_level'):
            result["education_level"] = job.education_level
        if hasattr(job, 'min_experience_years'):
            result["min_experience_years"] = job.min_experience_years
        if hasattr(job, 'max_experience_years'):
            result["max_experience_years"] = job.max_experience_years
        if hasattr(job, 'industry'):
            result["industry"] = job.industry
        if hasattr(job, 'tech_tags'):
            result["tech_tags"] = job.tech_tags
        if hasattr(job, 'company_name'):
            result["company_name"] = job.company_name
        if hasattr(job, 'company_type'):
            result["company_type"] = job.company_type
        if hasattr(job, 'company_size'):
            result["company_size"] = job.company_size
        if hasattr(job, 'company_intro'):
            result["company_intro"] = job.company_intro
        if hasattr(job, 'company_industry'):
            result["company_industry"] = job.company_industry
        if hasattr(job, 'company_tags'):
            result["company_tags"] = job.company_tags
        if hasattr(job, 'selected_benefits'):
            result["selected_benefits"] = job.selected_benefits

        return result

    def create_job(self, job_data: JobCreate, current_user: dict, db: Session) -> dict:
        self._init_repository(db)

        job_dict = job_data.dict()
        for key, value in job_dict.items():
            if isinstance(value, str):
                job_dict[key] = clean_special_chars(value)
        job_dict['user_id'] = current_user["user_id"]

        new_job = self.repository.create_job(job_dict)
        job_info = self._build_job_response(new_job)

        return {"message": "Job created successfully", "job": job_info}

    def get_jobs(self, current_user: dict, role: str, db: Session) -> list:
        self._init_repository(db)

        if role == "company":
            user_jobs = self.repository.get_jobs_by_user(current_user["user_id"])
        else:
            user_jobs = self.repository.get_all_jobs()

        return [self._build_job_response(job) for job in user_jobs]

    def get_job_detail(self, job_id: int, db: Session) -> dict:
        self._init_repository(db)
        
        job = self.repository.get_job_by_id(job_id)
        if not job:
            raise JobNotFoundException("Job not found")
        
        return self._build_job_response(job)

    def update_job(self, job_id: int, job_data: JobUpdate, current_user: dict, db: Session) -> dict:
        self._init_repository(db)
        
        job = self.repository.get_job_by_id(job_id)
        if not job:
            raise JobNotFoundException("Job not found")
        
        if job.user_id != current_user["user_id"]:
            raise AccessDeniedException("Access denied. You can only update your own jobs.")

        update_data = {}
        for key, value in job_data.dict().items():
            if value is not None:
                update_data[key] = clean_special_chars(value) if isinstance(value, str) else value

        updated_job = self.repository.update_job(job_id, update_data)
        job_info = self._build_job_response(updated_job)

        return {"message": "Job updated successfully", "job": job_info}

    def delete_job(self, job_id: int, current_user: dict, db: Session) -> dict:
        self._init_repository(db)
        
        job = self.repository.get_job_by_id(job_id)
        if not job:
            raise JobNotFoundException("Job not found")
        
        if job.user_id != current_user["user_id"]:
            raise AccessDeniedException("Access denied. You can only delete your own jobs.")

        self.repository.delete_job(job_id)
        return {"message": "Job deleted successfully"}

    async def upload_job_requirement(self, file: UploadFile, job_name: str, current_user: dict, db: Session, **kwargs) -> dict:
        self._init_repository(db)

        file_path = f"uploads/job_requirements/{file.filename}"
        os.makedirs(os.path.dirname(file_path), exist_ok=True)

        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        job_desc = ""
        if file.filename.endswith('.docx'):
            import docx
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                job_desc += paragraph.text + '\n'
        elif file.filename.endswith('.pdf'):
            try:
                import PyPDF2
                with open(file_path, 'rb') as pdf_file:
                    reader = PyPDF2.PdfReader(pdf_file)
                    for page_num in range(len(reader.pages)):
                        page = reader.pages[page_num]
                        job_desc += page.extract_text() + '\n'
            except ImportError:
                job_desc = "PDF解析需要安装PyPDF2库"
            except Exception as e:
                job_desc = f"PDF解析失败: {str(e)}"
        else:
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    job_desc = f.read()
            except:
                job_desc = "无法解析的文件格式"

        if len(job_desc) > 10000:
            job_desc = job_desc[:10000] + "...（内容过长，已截断）"

        job_data = {
            "job_name": clean_special_chars(job_name),
            "job_desc": clean_special_chars(job_desc),
            "salary": "",
            "location": "",
            "work_hours": "",
            "education_requirement": "",
            "experience_requirement": "",
            "recruitment_count": "",
            "department": "",
            "job_type": "",
            "benefits": "",
            "user_id": current_user["user_id"],
            "company_name": clean_special_chars(kwargs.get("company_name")) if kwargs.get("company_name") else None,
            "company_size": kwargs.get("company_size"),
            "company_industry": clean_special_chars(kwargs.get("company_industry")) if kwargs.get("company_industry") else None,
            "company_intro": clean_special_chars(kwargs.get("company_intro")) if kwargs.get("company_intro") else None
        }

        new_job = self.repository.create_job(job_data)
        job_info = {
            "id": new_job.id,
            "job_name": new_job.job_name,
            "job_desc": new_job.job_desc,
            "user_id": new_job.user_id,
            "create_time": new_job.created_at.strftime("%Y-%m-%d %H:%M:%S")
        }

        return {"message": "Job requirement uploaded successfully", "job": job_info}

    async def extract_job_info(self, files: list, current_user: dict) -> dict:
        full_text = ""
        for file in files:
            file_path = f"uploads/temp/{file.filename}"
            os.makedirs(os.path.dirname(file_path), exist_ok=True)

            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)

            if file.filename.endswith('.docx'):
                import docx
                doc = docx.Document(file_path)
                for paragraph in doc.paragraphs:
                    full_text += paragraph.text + '\n'
            elif file.filename.endswith('.pdf'):
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as pdf_file:
                        reader = PyPDF2.PdfReader(pdf_file)
                        for page_num in range(len(reader.pages)):
                            page = reader.pages[page_num]
                            full_text += page.extract_text() + '\n'
                except:
                    pass
            else:
                try:
                    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                        full_text += f.read() + '\n'
                except:
                    pass

        extracted_data = extract_job_info_from_text(full_text, files[0].filename if files else "")
        return extracted_data

    def create_job_ability_map(self, request_data: JobAbilityCreate, db: Session) -> dict:
        self._init_repository(db)
        
        job = self.repository.get_job_by_id(request_data.job_id)
        if not job:
            raise JobNotFoundException("Job not found")

        job_desc = job.job_desc or job.requirements or job.responsibilities or ""

        if not job_desc:
            ability_tree = {
                "name": job.job_name,
                "children": [
                    {
                        "name": "专业技能",
                        "children": [{"name": "编程语言"}, {"name": "技术框架"}, {"name": "数据库"}],
                        "weight": 0.4,
                        "match_type": "semantic"
                    },
                    {
                        "name": "工作经验",
                        "children": [{"name": "相关行业经验"}, {"name": "项目管理经验"}],
                        "weight": 0.3,
                        "match_type": "semantic"
                    },
                    {
                        "name": "学历要求",
                        "children": [{"name": "本科及以上"}],
                        "weight": 0.15,
                        "match_type": "exact"
                    },
                    {
                        "name": "证书资质",
                        "children": [],
                        "weight": 0.15,
                        "match_type": "semantic"
                    }
                ]
            }
        else:
            try:
                ability_report = _graph_agent.generate_ability_report(job_desc)
                ability_tree = {
                    "name": job.job_name,
                    "children": [
                        {
                            "name": "专业技能",
                            "children": [{"name": skill} for skill in ability_report.get("ability_tree", {}).get("children", [])[:5]],
                            "weight": 0.4,
                            "match_type": "semantic"
                        },
                        {
                            "name": "工作经验",
                            "children": [],
                            "weight": 0.3,
                            "match_type": "semantic"
                        },
                        {
                            "name": "学历要求",
                            "children": [],
                            "weight": 0.15,
                            "match_type": "exact"
                        },
                        {
                            "name": "证书资质",
                            "children": [],
                            "weight": 0.15,
                            "match_type": "semantic"
                        }
                    ]
                }
            except Exception:
                ability_tree = {
                    "name": job.job_name,
                    "children": [
                        {
                            "name": "专业技能",
                            "children": [{"name": "待分析"}],
                            "weight": 0.4,
                            "match_type": "semantic"
                        },
                        {
                            "name": "工作经验",
                            "children": [],
                            "weight": 0.3,
                            "match_type": "semantic"
                        },
                        {
                            "name": "学历要求",
                            "children": [],
                            "weight": 0.15,
                            "match_type": "exact"
                        },
                        {
                            "name": "证书资质",
                            "children": [],
                            "weight": 0.15,
                            "match_type": "semantic"
                        }
                    ]
                }

        return {"message": "岗位能力图谱创建成功", "data": ability_tree}

    def get_job_ability_tree(self, job_id: int, db: Session) -> dict:
        self._init_repository(db)
        
        job = self.repository.get_job_by_id(job_id)
        if not job:
            raise JobNotFoundException("Job not found")

        ability_tree = {
            "name": job.job_name,
            "children": [
                {
                    "name": "专业技能",
                    "children": [{"name": "Python"}, {"name": "Java"}, {"name": "SQL"}, {"name": "Git"}],
                    "weight": 0.4,
                    "match_type": "semantic"
                },
                {
                    "name": "工作经验",
                    "children": [],
                    "weight": 0.3,
                    "match_type": "semantic"
                },
                {
                    "name": "学历要求",
                    "children": [],
                    "weight": 0.15,
                    "match_type": "exact"
                },
                {
                    "name": "证书资质",
                    "children": [],
                    "weight": 0.15,
                    "match_type": "semantic"
                }
            ]
        }

        return ability_tree

    def update_job_ability_node(self, node_id: str, node_data: dict) -> dict:
        return {"message": "岗位能力节点更新成功", "data": node_data}

    def delete_job_ability_map(self, job_id: int) -> dict:
        return {"message": "岗位能力图谱删除成功"}