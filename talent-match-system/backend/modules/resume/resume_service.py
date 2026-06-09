import os
import json
from datetime import datetime
from modules.common.config import settings
from modules.resume.resume_repository import ResumeRepository
from modules.resume.schemas import ResumeUpdate
from modules.resume.exceptions import (
    InvalidFileTypeException,
    FileSizeExceededException,
    ResumeNotFoundException,
    AccessDeniedException
)
from modules.agents.resume_agent import ResumeParseAgent
from sqlalchemy.orm import Session
from fastapi import UploadFile


ALLOWED_EXTENSIONS = (".pdf", ".docx", ".jpg", ".jpeg", ".png", ".bmp", ".tiff")

# 全局Agent实例
_resume_agent = ResumeParseAgent()


def extract_resume_info(file_paths):
    """使用Agent解析简历文件（支持多格式）"""
    try:
        import docx
        from PyPDF2 import PdfReader
        
        resume_text = ""
        processed_count = 0
        
        for file_path in file_paths:
            if file_path.endswith('.docx'):
                doc = docx.Document(file_path)
                resume_text += "\n".join([para.text for para in doc.paragraphs])
                processed_count += 1
            elif file_path.endswith('.pdf'):
                with open(file_path, 'rb') as pdf_file:
                    reader = PdfReader(pdf_file)
                    for page in reader.pages:
                        text = page.extract_text()
                        if text:
                            resume_text += text + "\n"
                processed_count += 1
            elif file_path.lower().endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as txt_file:
                    resume_text += txt_file.read() + "\n"
                processed_count += 1
            elif file_path.lower().endswith(('.jpg', '.jpeg', '.png', '.bmp', '.tiff')):
                text = _extract_text_from_image(file_path)
                if text:
                    resume_text += text + "\n"
                    processed_count += 1
        
        if processed_count == 0:
            raise ValueError("未处理任何有效文件")
        
        result = _resume_agent.parse_resume(resume_text)
        
        work_exp = result.get("work_experience", [])
        internship_exp = result.get("internship_experience", [])
        edu_history = result.get("education_history", [])
        highlights = result.get("highlights", [])
        
        # 从工作经历和实习经历中提取项目经历
        projects = []
        
        # 从工作经历提取
        for exp in work_exp:
            if exp.get("description"):
                desc = exp["description"]
                if "项目" in desc or "负责" in desc or "参与" in desc:
                    projects.append({
                        "name": exp.get("position", "项目") + "项目",
                        "description": desc,
                        "role": exp.get("position", ""),
                        "skills": exp.get("technologies", []),
                        "domain": exp.get("domain", "")
                    })
        
        # 从实习经历提取项目和技术栈
        for exp in internship_exp:
            if exp.get("description"):
                desc = exp["description"]
                if "项目" in desc or "负责" in desc or "参与" in desc:
                    projects.append({
                        "name": "实习项目 - " + exp.get("position", "项目"),
                        "description": desc,
                        "role": exp.get("position", ""),
                        "skills": exp.get("technologies", []),
                        "domain": exp.get("domain", "")
                    })
        
        # 从highlights中提取
        for highlight in highlights:
            if "项目" in highlight:
                projects.append({
                    "name": "项目经历",
                    "description": highlight,
                    "role": "",
                    "skills": [],
                    "domain": ""
                })
        
        # 构建工作经历（正式工作）
        work_experience_list = []
        for exp in work_exp:
            work_experience_list.append({
                "period": exp.get("duration", ""),
                "company": exp.get("company", ""),
                "position": exp.get("position", ""),
                "description": exp.get("description", ""),
                "technologies": exp.get("technologies", []),
                "domain": exp.get("domain", "")
            })
        
        # 构建实习经历
        internship_list = []
        for exp in internship_exp:
            internship_list.append({
                "period": exp.get("duration", ""),
                "company": exp.get("company", ""),
                "position": exp.get("position", ""),
                "description": exp.get("description", ""),
                "technologies": exp.get("technologies", []),
                "domain": exp.get("domain", "")
            })
        
        # 获取工作经历的第一项（用于展示）
        first_work = work_exp[0] if work_exp else {}
        first_internship = internship_exp[0] if internship_exp else {}
        
        return {
            "personalInfo": {
                "name": result.get("name", "未知"),
                "phone": result.get("phone", ""),
                "email": result.get("email", ""),
                "address": ""
            },
            "education": {
                "period": edu_history[0]["duration"] if edu_history else "",
                "school": edu_history[0]["school"] if edu_history else "",
                "major": edu_history[0]["major"] if edu_history else "",
                "degree": result.get("education", "")
            },
            "workExperience": {
                "period": first_work.get("duration", ""),
                "company": first_work.get("company", ""),
                "position": first_work.get("position", ""),
                "description": first_work.get("description", ""),
                "technologies": first_work.get("technologies", []),
                "domain": first_work.get("domain", "")
            },
            "internshipExperience": {
                "period": first_internship.get("duration", ""),
                "company": first_internship.get("company", ""),
                "position": first_internship.get("position", ""),
                "description": first_internship.get("description", ""),
                "technologies": first_internship.get("technologies", []),
                "domain": first_internship.get("domain", "")
            },
            "skills": result.get("skills", []),
            "skill_tags": result.get("skill_tags", {}),
            "selfEvaluation": result.get("self_evaluation", ""),
            "projects": projects,
            "work_experience": work_experience_list,
            "internship_experience": internship_list,
            "education_history": edu_history,
            "_agent_result": result
        }
    except Exception as e:
        print(f"[ResumeService] Agent解析失败: {e}")
        return _extract_resume_fallback(file_paths)


def _extract_text_from_image(image_path: str) -> str:
    """使用PaddleOCR（国产）从图片中提取文字"""
    try:
        from paddleocr import PaddleOCR
        
        ocr = PaddleOCR(use_angle_cls=True, lang='ch', use_gpu=False, show_log=False)
        result = ocr.ocr(image_path, cls=True)
        
        if not result or not result[0]:
            return ""
        
        texts = []
        for line in result[0]:
            if line and len(line) >= 2:
                text = line[1][0]
                if text:
                    texts.append(text)
        
        return '\n'.join(texts)
    except ImportError:
        print("[ResumeService] 缺少PaddleOCR依赖，无法处理图片")
        return ""
    except Exception as e:
        print(f"[ResumeService] OCR识别失败: {e}")
        return ""


def _extract_resume_fallback(file_paths: list) -> dict:
    """从简历文件中提取信息（fallback方法）"""
    resume_text = ""
    
    try:
        import docx
        from PyPDF2 import PdfReader
        
        for file_path in file_paths:
            if file_path.endswith('.docx'):
                doc = docx.Document(file_path)
                resume_text += "\n".join([para.text for para in doc.paragraphs])
            elif file_path.endswith('.pdf'):
                with open(file_path, 'rb') as pdf_file:
                    reader = PdfReader(pdf_file)
                    for page in reader.pages:
                        text = page.extract_text()
                        if text:
                            resume_text += text + "\n"
            elif file_path.lower().endswith(('.jpg', '.jpeg', '.png')):
                resume_text += _extract_text_from_image(file_path)
    
    except Exception as e:
        print(f"[ResumeService] 解析文件失败: {e}")
    
    return _parse_resume_text_fallback(resume_text)


def _parse_resume_text_fallback(text: str) -> dict:
    """从文本中提取简历信息（fallback解析）"""
    import re
    
    result = {
        "personalInfo": {
            "name": "",
            "phone": "",
            "email": "",
            "address": ""
        },
        "education": {
            "period": "",
            "school": "",
            "major": "",
            "degree": ""
        },
        "workExperience": {
            "period": "",
            "company": "",
            "position": "",
            "description": ""
        },
        "skills": [],
        "selfEvaluation": "",
        "projects": [],
        "work_experience": [],
        "education_history": []
    }
    
    if not text:
        result["personalInfo"]["name"] = "未知"
        return result
    
    name_patterns = [r'姓名[\s：:]([\u4e00-\u9fa5]{2,4})', r'name[\s：:]([\u4e00-\u9fa5]{2,4})', r'([\u4e00-\u9fa5]{2,4})\s*简历']
    for pattern in name_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            result["personalInfo"]["name"] = match.group(1)
            break
    if not result["personalInfo"]["name"]:
        result["personalInfo"]["name"] = "未知"
    
    phone_match = re.search(r'1[3-9]\d{9}', text)
    if phone_match:
        result["personalInfo"]["phone"] = phone_match.group()
    
    email_match = re.search(r'[\w.-]+@[\w.-]+\.\w+', text)
    if email_match:
        result["personalInfo"]["email"] = email_match.group()
    
    edu_levels = ['博士', '硕士', '本科', '大专']
    for level in edu_levels:
        if level in text:
            result["education"]["degree"] = level
            break
    
    school_patterns = [r'毕业院校[\s：:]([^\n]+)', r'院校[\s：:]([^\n]+)', r'大学[\s：:]([^\n]+)', r'([\u4e00-\u9fa5]+大学)']
    for pattern in school_patterns:
        match = re.search(pattern, text)
        if match:
            result["education"]["school"] = match.group(1).strip()
            break
    
    major_patterns = [r'专业[\s：:]([^\n]+)', r'主修[\s：:]([^\n]+)']
    for pattern in major_patterns:
        match = re.search(pattern, text)
        if match:
            result["education"]["major"] = match.group(1).strip()
            break
    
    company_patterns = [r'公司[\s：:]([^\n]+)', r'工作单位[\s：:]([^\n]+)', r'任职于[\s：:]([^\n]+)']
    company = ""
    for pattern in company_patterns:
        match = re.search(pattern, text)
        if match:
            company = match.group(1).strip()
            result["workExperience"]["company"] = company
            break
    
    position_patterns = [r'职位[\s：:]([^\n]+)', r'岗位[\s：:]([^\n]+)', r'职务[\s：:]([^\n]+)']
    position = ""
    for pattern in position_patterns:
        match = re.search(pattern, text)
        if match:
            position = match.group(1).strip()
            result["workExperience"]["position"] = position
            break
    
    # 提取工作经历部分
    work_section = _extract_section(text, ['工作经历', '工作经验', '工作经验', '职业经历', '工作历史'])
    if work_section:
        result["workExperience"]["description"] = work_section[:500]
        # 添加到完整工作经历数组
        result["work_experience"].append({
            "company": company,
            "position": position,
            "duration": "",
            "description": work_section[:500]
        })
    
    # 提取项目经历
    project_section = _extract_section(text, ['项目经历', '项目经验', '项目经验', '项目'])
    projects = []
    if project_section:
        # 尝试从项目部分提取多个项目
        project_items = _split_into_projects(project_section)
        for i, item in enumerate(project_items):
            projects.append({
                "name": f"项目{i+1}",
                "description": item[:300],
                "role": position,
                "skills": []
            })
    result["projects"] = projects
    
    skill_keywords = ['python', 'java', 'javascript', 'vue', 'react', 'node', 'django', 'flask', 'mysql', 'postgresql', 'mongodb', 'redis', 'docker', 'kubernetes', 'html', 'css', 'git', 'linux', 'sql']
    text_lower = text.lower()
    for skill in skill_keywords:
        if skill in text_lower:
            result["skills"].append(skill.capitalize())
    result["skills"] = list(set(result["skills"]))
    
    eval_patterns = [r'自我评价[\s：:]([\s\S]*?)(?=\n\n|\Z)', r'个人简介[\s：:]([\s\S]*?)(?=\n\n|\Z)', r'自我描述[\s：:]([\s\S]*?)(?=\n\n|\Z)']
    for pattern in eval_patterns:
        match = re.search(pattern, text)
        if match:
            result["selfEvaluation"] = match.group(1).strip()[:200]
            break
    
    return result


def _extract_section(text: str, section_names: list) -> str:
    """提取指定部分的文本内容"""
    import re
    for name in section_names:
        pattern = rf'{name}[\s：:\-]*\s*([\s\S]*?)(?=\n\s*[^\s\n]{2,10}[\s：:\-]|\Z)'
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1).strip()
    return ""


def _split_into_projects(text: str) -> list:
    """将项目部分拆分为多个项目"""
    import re
    # 按常见分隔符拆分
    items = re.split(r'\n\d+\.|\n- |\n● |\n\* ', text)
    # 过滤短文本
    return [item.strip() for item in items if len(item.strip()) > 20]


import logging

logger = logging.getLogger(__name__)


class ResumeService:
    def __init__(self):
        self.repository = None
        os.makedirs(settings.UPLOAD_DIR, exist_ok=True)

    def _safe_parse_json(self, json_str, default=None):
        """安全解析JSON，失败时记录日志并返回默认值"""
        if not json_str:
            return default
        
        if isinstance(json_str, (list, dict)):
            return json_str
        
        try:
            return json.loads(json_str)
        except (json.JSONDecodeError, TypeError, ValueError) as e:
            logger.warning(f"JSON解析失败: {str(e)[:100]}, 原始数据: {str(json_str)[:50]}")
            return default

    def _normalize_field(self, value, expected_type=list):
        """确保字段类型正确，失败时返回默认值"""
        if value is None:
            return [] if expected_type == list else {}
        
        if isinstance(value, expected_type):
            return value
        
        if expected_type == list:
            if isinstance(value, str):
                try:
                    return json.loads(value)
                except:
                    return []
            elif isinstance(value, (tuple, set)):
                return list(value)
            return []
        elif expected_type == dict:
            if isinstance(value, str):
                try:
                    return json.loads(value)
                except:
                    return {}
            return {}
        
        return value

    def _init_repository(self, db: Session):
        if not self.repository:
            self.repository = ResumeRepository(db)

    def _validate_file(self, file: UploadFile):
        if not file.filename.lower().endswith(ALLOWED_EXTENSIONS):
            raise InvalidFileTypeException("Only PDF, DOCX and image files are allowed")

        file.file.seek(0, 2)
        file_size = file.file.tell()
        file.file.seek(0)

        if file_size > settings.MAX_UPLOAD_SIZE:
            raise FileSizeExceededException(f"File size exceeds the limit of {settings.MAX_UPLOAD_SIZE / 1024 / 1024}MB")

    def _build_resume_response(self, resume) -> dict:
        print(f"[DEBUG] _build_resume_response:")
        print(f"  resume.experiences: {resume.experiences}")
        print(f"  resume.honors: {resume.honors}")
        
        experiences = self._safe_parse_json(getattr(resume, "experiences", "[]"), [])
        honors = self._safe_parse_json(getattr(resume, "honors", "[]"), [])
        
        print(f"  解析后的experiences: {experiences}")
        print(f"  解析后的honors: {honors}")
        
        return {
            "id": resume.id,
            "resume_name": resume.resume_name,
            "filename": resume.filename,
            "filepath": resume.filepath,
            "user_id": resume.user_id,
            "name": resume.name,
            "phone": resume.phone,
            "email": resume.email,
            "desired_position": getattr(resume, "desired_position", ""),
            "address": resume.address,
            "education": resume.education,
            "school": resume.school,
            "major": resume.major,
            "degree": resume.degree,
            "skills": self._safe_parse_json(resume.skills, []),
            "skill_tags": self._safe_parse_json(getattr(resume, "skill_tags", "{}"), {}),
            "self_evaluation": resume.self_evaluation,
            "certificates": self._safe_parse_json(getattr(resume, "certificates", "[]"), []),
            "projects": self._safe_parse_json(getattr(resume, "projects", "[]"), []),
            "work_experience": self._safe_parse_json(getattr(resume, "work_experience", "[]"), []),
            "internship_experience": self._safe_parse_json(getattr(resume, "internship_experience", "[]"), []),
            "experiences": experiences,
            "honors": honors,
            "desired_salary": getattr(resume, "desired_salary", ""),
            "work_years": getattr(resume, "work_years", 0),
            "create_time": resume.created_at.strftime("%Y-%m-%d %H:%M:%S")
        }

    async def upload_resume(self, file: UploadFile, current_user: dict, db: Session, **kwargs) -> dict:
        self._init_repository(db)
        
        filename = "无文件"
        file_path = ""
        
        if file:
            self._validate_file(file)
            file_path = os.path.join(settings.UPLOAD_DIR, file.filename)
            try:
                with open(file_path, "wb") as f:
                    f.write(await file.read())
            except Exception as e:
                print(f"文件处理错误: {e}")
            filename = file.filename

        resume_data = {
            "resume_name": kwargs.get("resume_name"),
            "filename": filename,
            "filepath": file_path,
            "user_id": current_user["user_id"],
            "name": kwargs.get("name", "未知"),
            "age": kwargs.get("age", None),
            "phone": kwargs.get("phone", ""),
            "email": kwargs.get("email", ""),
            "desired_position": kwargs.get("desired_position", ""),
            "address": kwargs.get("address", ""),
            "education": kwargs.get("education", ""),
            "school": kwargs.get("school", ""),
            "major": kwargs.get("major", ""),
            "degree": kwargs.get("degree", "不限"),
            "graduation_date": kwargs.get("graduation_date", ""),
            "skills": kwargs.get("skills", "[]"),
            "self_evaluation": kwargs.get("self_evaluation", ""),
            "certificates": kwargs.get("certificates", "[]"),
            "projects": kwargs.get("projects", "[]"),
            "work_experience": kwargs.get("work_experience", "[]"),
            "internship_experience": kwargs.get("internship_experience", "[]"),
            "experiences": kwargs.get("experiences", "[]"),
            "honors": kwargs.get("honors", "[]"),
            "desired_salary": kwargs.get("desired_salary", ""),
            "work_years": kwargs.get("work_years", 0)
        }

        new_resume = self.repository.create_resume(resume_data)
        resume_info = self._build_resume_response(new_resume)

        return {"message": "Resume uploaded successfully", "resume": resume_info}

    def get_resumes(self, db: Session) -> list:
        self._init_repository(db)
        all_resumes = self.repository.get_all_resumes()
        return [self._build_resume_response(resume) for resume in all_resumes]

    def get_all_resumes_for_company(self, db: Session) -> list:
        self._init_repository(db)
        all_resumes = self.repository.get_all_resumes()
        return [self._build_resume_response(resume) for resume in all_resumes]

    def get_resume_detail(self, resume_id: int, db: Session) -> dict:
        self._init_repository(db)
        resume = self.repository.get_resume_by_id(resume_id)
        if not resume:
            raise ResumeNotFoundException("Resume not found")
        return self._build_resume_response(resume)

    def update_resume(self, resume_id: int, resume_data: ResumeUpdate, current_user: dict, db: Session) -> dict:
        self._init_repository(db)
        
        resume = self.repository.get_resume_by_id(resume_id)
        if not resume:
            raise ResumeNotFoundException("Resume not found")
        
        if resume.user_id != current_user["user_id"]:
            raise AccessDeniedException("You can only update your own resume")

        update_data = {
            "resume_name": resume_data.resume_name,
            "name": resume_data.name or "未知",
            "age": getattr(resume_data, "age", None),
            "phone": resume_data.phone or "",
            "email": resume_data.email or "",
            "desired_position": resume_data.desired_position or "",
            "address": resume_data.address or "",
            "education": resume_data.education or "",
            "school": resume_data.school or "",
            "major": resume_data.major or "",
            "degree": resume_data.degree or "不限",
            "graduation_date": getattr(resume_data, "graduation_date", ""),
            "skills": resume_data.skills or "[]",
            "self_evaluation": resume_data.self_evaluation or "",
            "certificates": resume_data.certificates or "[]",
            "projects": resume_data.projects or "[]",
            "work_experience": resume_data.work_experience or "[]",
            "internship_experience": resume_data.internship_experience or "[]",
            "experiences": getattr(resume_data, "experiences", "[]"),
            "honors": getattr(resume_data, "honors", "[]"),
            "desired_salary": resume_data.desired_salary or "",
            "work_years": resume_data.work_years or 0
        }

        updated_resume = self.repository.update_resume(resume_id, update_data)
        result = self._build_resume_response(updated_resume)

        return {"message": "Resume updated successfully", "resume": result}

    def delete_resume(self, resume_id: int, current_user: dict, db: Session) -> dict:
        self._init_repository(db)
        
        resume = self.repository.get_resume_by_id(resume_id)
        if not resume:
            raise ResumeNotFoundException("Resume not found")
        
        if resume.user_id != current_user["user_id"]:
            raise AccessDeniedException("You can only delete your own resume")

        success = self.repository.delete_resume(resume_id)
        if not success:
            raise Exception("Failed to delete resume")
        
        return {"message": "Resume deleted successfully"}

    async def extract_resume_info_endpoint(self, files: list) -> dict:
        temp_file_paths = []

        for file in files:
            if file.filename.endswith(".pdf") or file.filename.endswith(".docx") or file.filename.lower().endswith(".txt") or file.filename.lower().endswith(('.jpg', '.jpeg', '.png', '.bmp', '.tiff')):
                temp_file_path = os.path.join(settings.UPLOAD_DIR, f"temp_{datetime.now().timestamp()}_{file.filename}")
                try:
                    with open(temp_file_path, "wb") as f:
                        f.write(await file.read())
                    temp_file_paths.append(temp_file_path)
                except Exception as e:
                    print(f"文件保存错误: {e}")
                    continue

        try:
            extracted_info = extract_resume_info(temp_file_paths)
        except Exception as e:
            print(f"AI解析错误: {e}")
            # 不返回模拟数据，而是使用 fallback 解析
            try:
                extracted_info = _extract_resume_fallback(temp_file_paths)
            except Exception as e2:
                print(f"Fallback解析也失败: {e2}")
                # 返回空结果
                extracted_info = {
                    "personalInfo": {
                        "name": "未知",
                        "phone": "",
                        "email": "",
                        "address": ""
                    },
                    "education": {
                        "period": "",
                        "school": "",
                        "major": "",
                        "degree": ""
                    },
                    "workExperience": {
                        "period": "",
                        "company": "",
                        "position": "",
                        "description": ""
                    },
                    "skills": [],
                    "selfEvaluation": "",
                    "projects": [],
                    "work_experience": [],
                    "education_history": []
                }
        finally:
            for temp_file_path in temp_file_paths:
                if os.path.exists(temp_file_path):
                    os.remove(temp_file_path)

        return extracted_info